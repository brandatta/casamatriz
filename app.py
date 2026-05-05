import base64
import io
import re
from datetime import datetime
from pathlib import Path
from urllib.parse import quote

import pandas as pd
import requests
import streamlit as st
from docx import Document
from pypdf import PdfReader


# ============================================================
# CONFIG
# ============================================================

st.set_page_config(
    page_title="Casa Matriz | Archivo organizado",
    page_icon="🗂️",
    layout="wide",
)

CATEGORIES = [
    "01_IDENTIDAD_MARCA_Y_ESTRATEGIA",
    "02_OBRA_AUTORAL_ENSAYOS_Y_BORRADORES",
    "03_CURSOS_FORMACION_Y_ARCHIVO_ASTROLOGICO",
    "04_BIBLIOGRAFIA_INVESTIGACION_Y_FUENTES",
    "05_ARCHIVO_VISUAL_EDITORIAL_Y_REFERENCIAS",
]

CATEGORY_LABELS = {
    "01_IDENTIDAD_MARCA_Y_ESTRATEGIA": "Identidad, marca y estrategia",
    "02_OBRA_AUTORAL_ENSAYOS_Y_BORRADORES": "Obra autoral, ensayos y borradores",
    "03_CURSOS_FORMACION_Y_ARCHIVO_ASTROLOGICO": "Cursos, formación y archivo astrológico",
    "04_BIBLIOGRAFIA_INVESTIGACION_Y_FUENTES": "Bibliografía, investigación y fuentes",
    "05_ARCHIVO_VISUAL_EDITORIAL_Y_REFERENCIAS": "Archivo visual, editorial y referencias",
}

CATEGORY_DESCRIPTIONS = {
    "01_IDENTIDAD_MARCA_Y_ESTRATEGIA": "Material fundacional, marca, manifiestos, nombres, estructura y posicionamiento.",
    "02_OBRA_AUTORAL_ENSAYOS_Y_BORRADORES": "Textos propios, ensayos, fragmentos, ideas y materiales publicables.",
    "03_CURSOS_FORMACION_Y_ARCHIVO_ASTROLOGICO": "Cursos, certificaciones, seminarios, material pedagógico y archivo histórico astrológico.",
    "04_BIBLIOGRAFIA_INVESTIGACION_Y_FUENTES": "Tesis, PDFs académicos, libros de referencia, papers y bibliografía.",
    "05_ARCHIVO_VISUAL_EDITORIAL_Y_REFERENCIAS": "Imágenes, recursos visuales, links, referencias editoriales y diseño.",
}

NAV_ITEMS = {
    "🏠 Inicio / Subir archivos": "home",
    "📁 Identidad, marca y estrategia": "01_IDENTIDAD_MARCA_Y_ESTRATEGIA",
    "📁 Obra autoral, ensayos y borradores": "02_OBRA_AUTORAL_ENSAYOS_Y_BORRADORES",
    "📁 Cursos, formación y archivo astrológico": "03_CURSOS_FORMACION_Y_ARCHIVO_ASTROLOGICO",
    "📁 Bibliografía, investigación y fuentes": "04_BIBLIOGRAFIA_INVESTIGACION_Y_FUENTES",
    "📁 Archivo visual, editorial y referencias": "05_ARCHIVO_VISUAL_EDITORIAL_Y_REFERENCIAS",
    "📋 Inventario general": "inventory",
}

IMAGE_EXTENSIONS = {
    ".jpg", ".jpeg", ".png", ".webp", ".gif", ".tif", ".tiff", ".jp2"
}

DESIGN_EXTENSIONS = {
    ".ai", ".psd", ".indd"
}

RULES = {
    "01_IDENTIDAD_MARCA_Y_ESTRATEGIA": [
        "casa matriz",
        "proyecto integral",
        "manifiesto",
        "fundamentos",
        "marca",
        "identidad",
        "estrategia",
        "nombres",
        "tagline",
        "estructura",
        "talleres",
        "publicaciones",
        "matrices",
        "mediatrices",
        "editable marcas",
        "landing",
        "web",
    ],
    "02_OBRA_AUTORAL_ENSAYOS_Y_BORRADORES": [
        "ensayo",
        "cosmos",
        "core",
        "pliegues",
        "habitar",
        "venus",
        "buena astrologa",
        "ideas",
        "otros titulos",
        "varios",
        "adorno",
        "foucault",
        "hermeneutica",
        "mujeres",
        "giro copernicano",
        "fragmento",
        "borrador",
        "dossier",
        "cuerpo",
        "metamorfosis",
        "oscuridad",
        "sirenas",
        "criptozoologia",
    ],
    "03_CURSOS_FORMACION_Y_ARCHIVO_ASTROLOGICO": [
        "curso",
        "formacion",
        "certificacion",
        "luminarias",
        "sol y luna",
        "intro curso",
        "preview",
        "material adicional",
        "seminario",
        "volver a la luna",
        "transcripcion",
        "horoscono",
        "ernesto castro",
        "arquetipo",
        "lunar",
        "solar",
        "astrologia",
        "carta natal",
        "zodiaco",
        "tauro",
        "libra",
        "luna astrologica",
    ],
    "04_BIBLIOGRAFIA_INVESTIGACION_Y_FUENTES": [
        "bibliografia",
        "tesis",
        "thesis",
        "phd",
        "startup",
        "phillipson",
        "clynes",
        "von stuckrad",
        "astrology and truth",
        "internet on modern western astrology",
        "historia",
        "borges",
        "seres imaginarios",
        "libro de los seres",
        "fuente",
        "paper",
        "research",
        "epistemology",
        "references",
        "bibliography",
        "university",
        "submitted",
        "abstract",
        "table of contents",
        "libro",
        "book",
        "journal",
        "chapter",
    ],
    "05_ARCHIVO_VISUAL_EDITORIAL_Y_REFERENCIAS": [
        "cabinet",
        "wonders",
        "collection",
        "coleccion ilustrada",
        "links",
        "rawpixel",
        "dragon",
        "utagawa",
        "kuniyoshi",
        "imagen",
        "visual",
        "referencia",
        "archivo visual",
        "public domain",
        "wellcome",
        "british library",
        "library of congress",
        "grabado",
        "manuscrito",
        "sketchbook",
        "stephan scriber",
        "ilustrada",
        "ilustrado",
        "editorial",
        "wunderkammer",
        "kunstkammer",
    ],
}


# ============================================================
# GITHUB CONFIG
# ============================================================

def get_github_config():
    token = st.secrets.get("GITHUB_TOKEN", "")
    repo = st.secrets.get("GITHUB_REPO", "")
    branch = st.secrets.get("GITHUB_BRANCH", "main")

    if not token or not repo:
        st.error(
            "Faltan secrets de GitHub. Configurá GITHUB_TOKEN, GITHUB_REPO y GITHUB_BRANCH en Streamlit Cloud."
        )
        st.stop()

    return token, repo, branch


def github_headers():
    token, _, _ = get_github_config()
    return {
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    }


def github_api_url(path: str) -> str:
    _, repo, _ = get_github_config()
    encoded_path = quote(path, safe="")
    return f"https://api.github.com/repos/{repo}/contents/{encoded_path}"


def github_get_file(path: str):
    _, _, branch = get_github_config()

    response = requests.get(
        github_api_url(path),
        headers=github_headers(),
        params={"ref": branch},
        timeout=30,
    )

    if response.status_code == 404:
        return None

    response.raise_for_status()
    return response.json()


def github_list_dir(path: str):
    _, _, branch = get_github_config()

    response = requests.get(
        github_api_url(path),
        headers=github_headers(),
        params={"ref": branch},
        timeout=30,
    )

    if response.status_code == 404:
        return []

    response.raise_for_status()
    data = response.json()

    if isinstance(data, list):
        return data

    return []


def github_put_file(path: str, file_bytes: bytes, message: str):
    _, _, branch = get_github_config()

    existing = github_get_file(path)

    payload = {
        "message": message,
        "content": base64.b64encode(file_bytes).decode("utf-8"),
        "branch": branch,
    }

    if existing and existing.get("sha"):
        payload["sha"] = existing["sha"]

    response = requests.put(
        github_api_url(path),
        headers=github_headers(),
        json=payload,
        timeout=90,
    )

    response.raise_for_status()
    return response.json()


def github_delete_file(path: str, message: str):
    _, _, branch = get_github_config()

    existing = github_get_file(path)

    if not existing:
        return None

    payload = {
        "message": message,
        "sha": existing["sha"],
        "branch": branch,
    }

    response = requests.delete(
        github_api_url(path),
        headers=github_headers(),
        json=payload,
        timeout=60,
    )

    response.raise_for_status()
    return response.json()


def github_download_file(path: str) -> bytes:
    """
    Descarga archivos desde GitHub.
    Soporta:
    1. Archivos chicos devueltos como base64 por Contents API.
    2. Archivos grandes con download_url.
    3. Archivos grandes/privados usando Accept: application/vnd.github.raw.
    """
    data = github_get_file(path)

    if not data:
        raise FileNotFoundError(path)

    content = data.get("content")
    encoding = data.get("encoding")
    download_url = data.get("download_url")

    # Caso 1: archivo chico devuelto en base64
    if content and encoding == "base64":
        try:
            return base64.b64decode(content)
        except Exception:
            pass

    # Caso 2: archivo con download_url
    if download_url:
        try:
            response = requests.get(
                download_url,
                headers=github_headers(),
                timeout=120,
            )
            response.raise_for_status()

            if response.content:
                return response.content
        except Exception:
            pass

    # Caso 3: fallback raw vía GitHub Contents API
    # Esto suele resolver PDFs grandes o repos privados.
    _, _, branch = get_github_config()

    raw_headers = github_headers()
    raw_headers["Accept"] = "application/vnd.github.raw"

    response = requests.get(
        github_api_url(path),
        headers=raw_headers,
        params={"ref": branch},
        timeout=120,
    )

    response.raise_for_status()

    if response.content:
        return response.content

    raise ValueError(f"No se pudo descargar {path}")

def github_move_file(old_path: str, new_path: str, file_bytes: bytes, filename: str):
    """
    GitHub Contents API no tiene move nativo.
    Se hace put en nuevo path y delete en path viejo.
    """
    if old_path == new_path:
        return

    github_put_file(
        new_path,
        file_bytes,
        f"Mover {filename} a {new_path}",
    )

    github_delete_file(
        old_path,
        f"Eliminar ubicación anterior de {filename}",
    )


# ============================================================
# TEXT / CLASSIFICATION HELPERS
# ============================================================

def normalize(value: str) -> str:
    if not value:
        return ""

    value = value.lower()
    value = value.replace("_", " ").replace("-", " ")
    value = re.sub(r"\s+", " ", value)

    replacements = {
        "á": "a", "é": "e", "í": "i", "ó": "o",
        "ú": "u", "ü": "u", "ñ": "n",
    }

    for original, replacement in replacements.items():
        value = value.replace(original, replacement)

    return value.strip()


def clean_filename(filename: str) -> str:
    filename = filename.strip()
    filename = re.sub(r'[<>:"/\\|?*]', "-", filename)
    filename = re.sub(r"\s+", " ", filename)
    return filename[:180]


def extract_docx(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
        parts = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    except Exception as error:
        return f"[ERROR_DOCX] {error}"


def extract_pdf(file_bytes: bytes, max_pages: int = 5) -> str:
    """
    Se mantiene para tags/metadatos si alguna vez se usa.
    La clasificación de PDF es absoluta a Bibliografía.
    """
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = []

        total_pages = min(len(reader.pages), max_pages)

        for index in range(total_pages):
            text = reader.pages[index].extract_text() or ""
            if text.strip():
                parts.append(text.strip())

        return "\n".join(parts)

    except Exception as error:
        return f"[ERROR_PDF] {error}"


def extract_txt(file_bytes: bytes) -> str:
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue

    return "[ERROR_TXT] No se pudo leer el archivo de texto."


def extract_text(filename: str, file_bytes: bytes) -> str:
    extension = Path(filename).suffix.lower()

    if extension == ".docx":
        return extract_docx(file_bytes)

    if extension == ".pdf":
        return extract_pdf(file_bytes)

    if extension in [".txt", ".md"]:
        return extract_txt(file_bytes)

    if extension in IMAGE_EXTENSIONS:
        return "[IMAGEN] Archivo visual."

    if extension in DESIGN_EXTENSIONS:
        return "[DISEÑO] Archivo de diseño."

    return "[SIN_TEXTO] No se extrajo texto para este tipo de archivo."


def word_count(text: str) -> int:
    if not text or text.startswith("["):
        return 0

    return len(re.findall(r"\b\w+\b", text))


def detect_tags(text: str) -> str:
    normalized = normalize(text)

    tag_rules = {
        "astrologia": ["astrologia", "zodiaco", "carta natal", "horoscopo"],
        "cine": ["cine", "pelicula", "pantalla", "fotograma"],
        "luna": ["luna", "lunar", "selene", "artemisa"],
        "sol": ["sol", "solar"],
        "mitologia": ["mitologia", "dioses", "diosa", "arquetipo"],
        "bestiario": ["bestiario", "grifo", "fenix", "salamandra", "unicornio", "dragon"],
        "imagen": ["imagen", "visual", "grabado", "manuscrito", "ilustracion"],
        "marca": ["marca", "identidad", "logo", "tagline"],
        "curso": ["curso", "clase", "certificacion", "material adicional"],
        "bibliografia": ["bibliografia", "references", "bibliography", "tesis", "thesis", "paper"],
        "tesis": ["tesis", "thesis", "phd"],
        "ensayo": ["ensayo", "dossier", "texto"],
        "editorial": ["editorial", "coleccion", "cabinet", "wonders"],
    }

    tags = []

    for tag, keywords in tag_rules.items():
        for keyword in keywords:
            if keyword in normalized:
                tags.append(tag)
                break

    return ", ".join(sorted(set(tags)))


def classify_file(filename: str, text: str):
    extension = Path(filename).suffix.lower()

    # ------------------------------------------------------------
    # REGLA ABSOLUTA
    # ------------------------------------------------------------
    # Todo PDF va siempre a Bibliografía / investigación / fuentes.
    # No se evalúan keywords ni contenido para PDFs.
    if extension == ".pdf":
        return (
            "04_BIBLIOGRAFIA_INVESTIGACION_Y_FUENTES",
            999,
            "PDF clasificado automáticamente como bibliografía/fuente",
        )

    normalized_name = normalize(filename)
    normalized_text = normalize(text[:6000])

    scores = {category: 0 for category in CATEGORIES}
    reasons = {category: [] for category in CATEGORIES}

    # ------------------------------------------------------------
    # DEFAULTS FUERTES POR EXTENSIÓN
    # ------------------------------------------------------------

    if extension in IMAGE_EXTENSIONS:
        scores["05_ARCHIVO_VISUAL_EDITORIAL_Y_REFERENCIAS"] += 30
        reasons["05_ARCHIVO_VISUAL_EDITORIAL_Y_REFERENCIAS"].append(
            f"archivo visual {extension}"
        )

    if extension in DESIGN_EXTENSIONS:
        scores["05_ARCHIVO_VISUAL_EDITORIAL_Y_REFERENCIAS"] += 30
        reasons["05_ARCHIVO_VISUAL_EDITORIAL_Y_REFERENCIAS"].append(
            f"archivo de diseño {extension}"
        )

    # ------------------------------------------------------------
    # KEYWORDS GENERALES PARA NO-PDF
    # ------------------------------------------------------------

    for category, keywords in RULES.items():
        for keyword in keywords:
            keyword_normalized = normalize(keyword)

            if keyword_normalized in normalized_name:
                scores[category] += 8
                reasons[category].append(f"nombre contiene '{keyword}'")

            elif keyword_normalized in normalized_text:
                scores[category] += 3
                reasons[category].append(f"contenido contiene '{keyword}'")

    # ------------------------------------------------------------
    # SELECCIÓN FINAL
    # ------------------------------------------------------------

    best_category = max(scores, key=scores.get)
    best_score = scores[best_category]

    if best_score == 0:
        return (
            "02_OBRA_AUTORAL_ENSAYOS_Y_BORRADORES",
            0,
            "sin coincidencias fuertes; revisar manualmente",
        )

    return best_category, best_score, "; ".join(reasons[best_category][:4])


def file_icon(extension: str) -> str:
    if extension in IMAGE_EXTENSIONS:
        return "🖼️"
    if extension == ".pdf":
        return "📕"
    if extension == ".docx":
        return "📄"
    if extension in [".txt", ".md"]:
        return "📝"
    if extension in DESIGN_EXTENSIONS:
        return "🎨"
    return "📎"


# ============================================================
# INVENTORY
# ============================================================

INVENTORY_PATH = "data/inventory.csv"


def empty_inventory() -> pd.DataFrame:
    return pd.DataFrame(
        columns=[
            "archivo",
            "extension",
            "tamano_kb",
            "categoria",
            "path",
            "score",
            "tags",
            "motivo",
            "palabras_extraidas",
            "uploaded_at",
        ]
    )


def load_inventory() -> pd.DataFrame:
    try:
        file_bytes = github_download_file(INVENTORY_PATH)
        df = pd.read_csv(io.BytesIO(file_bytes))

        for col in empty_inventory().columns:
            if col not in df.columns:
                df[col] = ""

        return df

    except Exception:
        return empty_inventory()


def save_inventory(df: pd.DataFrame):
    csv_bytes = df.to_csv(index=False).encode("utf-8-sig")
    github_put_file(
        INVENTORY_PATH,
        csv_bytes,
        "Actualizar inventario Casa Matriz",
    )


def add_inventory_row(row: dict):
    df = load_inventory()

    if not df.empty and "path" in df.columns:
        df = df[df["path"] != row["path"]]

    df = pd.concat([df, pd.DataFrame([row])], ignore_index=True)

    save_inventory(df)


def remove_inventory_path(path: str):
    df = load_inventory()

    if df.empty:
        return

    df = df[df["path"] != path]

    save_inventory(df)


def rebuild_inventory_from_storage() -> pd.DataFrame:
    """
    Reconstruye el inventario desde las carpetas reales de GitHub.
    Ojo: esta función NO mueve archivos.
    Pero si encuentra un PDF en carpeta incorrecta, lo marca como Bibliografía en el inventario.
    Para mover físicamente PDFs, usar force_move_all_pdfs_to_bibliography().
    """
    rows = []

    for category in CATEGORIES:
        contents = github_list_dir(f"storage/{category}")

        for item in contents:
            if item.get("type") != "file":
                continue

            filename = item["name"]
            extension = Path(filename).suffix.lower()
            path = item["path"]

            final_category = (
                "04_BIBLIOGRAFIA_INVESTIGACION_Y_FUENTES"
                if extension == ".pdf"
                else category
            )

            rows.append(
                {
                    "archivo": filename,
                    "extension": extension,
                    "tamano_kb": round(item.get("size", 0) / 1024, 2),
                    "categoria": final_category,
                    "path": path,
                    "score": 999 if extension == ".pdf" else "",
                    "tags": "bibliografia" if extension == ".pdf" else "",
                    "motivo": (
                        "PDF clasificado automáticamente como bibliografía/fuente"
                        if extension == ".pdf"
                        else "detectado desde GitHub storage"
                    ),
                    "palabras_extraidas": "",
                    "uploaded_at": "",
                }
            )

    if not rows:
        return empty_inventory()

    return pd.DataFrame(rows)


def reclassify_and_move_existing_files() -> tuple[pd.DataFrame, list[str]]:
    """
    Lee el inventario actual, descarga cada archivo, reclasifica con las reglas actuales,
    mueve el archivo si cambió de categoría y actualiza el inventario.
    """
    df = load_inventory()

    if df.empty:
        return df, ["No hay archivos para reclasificar."]

    updated_rows = []
    logs = []

    for _, row in df.iterrows():
        old_path = str(row.get("path", ""))
        filename = str(row.get("archivo", ""))

        if not old_path or not filename:
            logs.append(f"SKIP: fila inválida sin path o archivo: {row.to_dict()}")
            continue

        try:
            file_bytes = github_download_file(old_path)
        except Exception as error:
            logs.append(f"ERROR descargando {filename}: {error}")
            updated_rows.append(row.to_dict())
            continue

        extension = Path(filename).suffix.lower()

        text = extract_text(filename, file_bytes)
        new_category, score, reason = classify_file(filename, text)
        tags = detect_tags(f"{filename} {text[:4000]}")

        old_category = str(row.get("categoria", ""))
        new_path = f"storage/{new_category}/{filename}"

        if old_path != new_path:
            try:
                github_move_file(
                    old_path=old_path,
                    new_path=new_path,
                    file_bytes=file_bytes,
                    filename=filename,
                )
                logs.append(f"MOVIDO: {filename} | {old_category} → {new_category}")
            except Exception as error:
                logs.append(f"ERROR moviendo {filename}: {error}")
                updated_rows.append(row.to_dict())
                continue
        else:
            logs.append(f"OK: {filename} sigue en {new_category}")

        updated_rows.append(
            {
                "archivo": filename,
                "extension": extension,
                "tamano_kb": round(len(file_bytes) / 1024, 2),
                "categoria": new_category,
                "path": new_path,
                "score": score,
                "tags": tags,
                "motivo": reason,
                "palabras_extraidas": word_count(text),
                "uploaded_at": row.get("uploaded_at", ""),
            }
        )

    new_df = pd.DataFrame(updated_rows)

    for col in empty_inventory().columns:
        if col not in new_df.columns:
            new_df[col] = ""

    new_df = new_df[empty_inventory().columns]

    save_inventory(new_df)

    return new_df, logs


def force_move_all_pdfs_to_bibliography() -> tuple[pd.DataFrame, list[str]]:
    """
    Escanea directamente todas las carpetas storage/* en GitHub.
    Si encuentra PDFs fuera de Bibliografía, los mueve físicamente.
    Luego reconstruye y guarda el inventario.
    No depende del inventario actual.
    """
    target_category = "04_BIBLIOGRAFIA_INVESTIGACION_Y_FUENTES"
    logs = []

    for category in CATEGORIES:
        contents = github_list_dir(f"storage/{category}")

        for item in contents:
            if item.get("type") != "file":
                continue

            filename = item["name"]
            extension = Path(filename).suffix.lower()
            old_path = item["path"]

            if extension != ".pdf":
                continue

            new_path = f"storage/{target_category}/{filename}"

            if old_path == new_path:
                logs.append(f"OK: {filename} ya está en Bibliografía")
                continue

            try:
                file_bytes = github_download_file(old_path)

                github_move_file(
                    old_path=old_path,
                    new_path=new_path,
                    file_bytes=file_bytes,
                    filename=filename,
                )

                logs.append(f"MOVIDO PDF: {filename} | {category} → {target_category}")

            except Exception as error:
                logs.append(f"ERROR moviendo {filename}: {error}")

    rebuilt = rebuild_inventory_from_storage()
    save_inventory(rebuilt)

    return rebuilt, logs


# ============================================================
# UI HELPERS
# ============================================================

def render_file_card(row):
    extension = str(row.get("extension", ""))
    icon = file_icon(extension)
    filename = str(row.get("archivo", ""))
    path = str(row.get("path", ""))

    with st.container(border=True):
        st.markdown(f"### {icon} {filename}")

        st.caption(
            f"{extension} · {row.get('tamano_kb', '')} KB · "
            f"{row.get('palabras_extraidas', '')} palabras"
        )

        if str(row.get("tags", "")).strip():
            st.markdown(f"**Tags:** `{row.get('tags')}`")

        if str(row.get("motivo", "")).strip():
            st.markdown(f"**Motivo:** {row.get('motivo')}")

        if str(row.get("uploaded_at", "")).strip():
            st.caption(f"Subido: {row.get('uploaded_at')}")

        try:
            file_bytes = github_download_file(path)

            st.download_button(
                "Descargar",
                data=file_bytes,
                file_name=filename,
                key=f"download_{path}",
            )
        except Exception as error:
            st.warning(f"No se pudo preparar la descarga: {error}")

        if st.button("Eliminar del sitio", key=f"delete_{path}"):
            try:
                github_delete_file(path, f"Eliminar {filename}")
                remove_inventory_path(path)
                st.success("Archivo eliminado.")
                st.rerun()
            except Exception as error:
                st.error(f"No se pudo eliminar: {error}")


# ============================================================
# SIDEBAR
# ============================================================

st.sidebar.title("🗂️ Casa Matriz")
st.sidebar.caption("Archivo persistente en GitHub")

selected_nav_label = st.sidebar.radio(
    "Navegación",
    list(NAV_ITEMS.keys()),
)

selected_page = NAV_ITEMS[selected_nav_label]

st.sidebar.divider()

inventory_df = load_inventory()

for category in CATEGORIES:
    count = int((inventory_df["categoria"] == category).sum()) if not inventory_df.empty else 0
    st.sidebar.caption(f"{CATEGORY_LABELS[category]}: {count}")


# ============================================================
# HOME
# ============================================================

if selected_page == "home":
    st.title("🗂️ Casa Matriz | Archivo organizado")

    st.write(
        "Subí archivos y la app los guarda permanentemente en GitHub, "
        "dentro de su categoría correspondiente."
    )

    st.info(
        "Regla actual: todo PDF va siempre a Bibliografía / investigación / fuentes."
    )

    uploaded_files = st.file_uploader(
        "Subir archivos",
        accept_multiple_files=True,
        type=[
            "docx", "pdf", "txt", "md",
            "jpg", "jpeg", "png", "webp", "gif", "tif", "tiff", "jp2",
            "ai", "psd", "indd",
        ],
    )

    if uploaded_files:
        if st.button("Guardar archivos en el sitio", type="primary"):
            with st.spinner("Clasificando y guardando en GitHub..."):
                for uploaded_file in uploaded_files:
                    filename = clean_filename(uploaded_file.name)
                    extension = Path(filename).suffix.lower()
                    file_bytes = uploaded_file.getvalue()

                    text = extract_text(filename, file_bytes)
                    category, score, reason = classify_file(filename, text)
                    tags = detect_tags(f"{filename} {text[:4000]}")

                    github_path = f"storage/{category}/{filename}"

                    github_put_file(
                        github_path,
                        file_bytes,
                        f"Subir archivo {filename} a {category}",
                    )

                    row = {
                        "archivo": filename,
                        "extension": extension,
                        "tamano_kb": round(len(file_bytes) / 1024, 2),
                        "categoria": category,
                        "path": github_path,
                        "score": score,
                        "tags": tags,
                        "motivo": reason,
                        "palabras_extraidas": word_count(text),
                        "uploaded_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    }

                    add_inventory_row(row)

            st.success("Archivos guardados en el sitio.")
            st.rerun()

    st.subheader("Secciones")

    inventory_df = load_inventory()

    cols = st.columns(5)

    for idx, category in enumerate(CATEGORIES):
        with cols[idx]:
            count = int((inventory_df["categoria"] == category).sum()) if not inventory_df.empty else 0
            st.metric(CATEGORY_LABELS[category], count)
            st.caption(CATEGORY_DESCRIPTIONS[category])

    st.stop()


# ============================================================
# CATEGORY PAGES
# ============================================================

if selected_page in CATEGORIES:
    category = selected_page

    st.title(f"📁 {CATEGORY_LABELS[category]}")
    st.caption(CATEGORY_DESCRIPTIONS[category])

    inventory_df = load_inventory()

    if inventory_df.empty:
        st.info("Todavía no hay archivos guardados.")
        st.stop()

    category_df = inventory_df[inventory_df["categoria"] == category].copy()

    st.metric("Archivos en esta sección", len(category_df))

    if category_df.empty:
        st.info("Todavía no hay archivos en esta sección.")
        st.stop()

    search = st.text_input("Buscar dentro de esta sección", "")

    if search:
        search_norm = normalize(search)
        category_df = category_df[
            category_df.apply(
                lambda row: search_norm in normalize(str(row.get("archivo", "")))
                or search_norm in normalize(str(row.get("tags", "")))
                or search_norm in normalize(str(row.get("motivo", ""))),
                axis=1,
            )
        ]

    cols = st.columns(3)

    for index, row in category_df.reset_index(drop=True).iterrows():
        with cols[index % 3]:
            render_file_card(row)

    st.stop()


# ============================================================
# INVENTORY
# ============================================================

if selected_page == "inventory":
    st.title("📋 Inventario general")

    st.warning(
        "Si todavía ves PDFs en secciones incorrectas, usá primero "
        "'Forzar PDFs a Bibliografía'. Ese botón escanea las carpetas reales de GitHub "
        "y no depende del inventario."
    )

    col_a, col_b, col_c = st.columns([1, 1, 1])

    with col_a:
        if st.button("Reconstruir inventario desde GitHub"):
            with st.spinner("Leyendo storage desde GitHub..."):
                rebuilt = rebuild_inventory_from_storage()
                save_inventory(rebuilt)
                st.success("Inventario reconstruido desde carpetas actuales.")
                st.rerun()

    with col_b:
        if st.button("Reclasificar y mover archivos existentes", type="primary"):
            with st.spinner("Reclasificando archivos y moviendo carpetas en GitHub..."):
                new_df, logs = reclassify_and_move_existing_files()
                st.success("Reclasificación finalizada.")
                with st.expander("Ver log de reclasificación", expanded=True):
                    for line in logs:
                        st.write(line)
                st.rerun()

    with col_c:
        if st.button("Forzar PDFs a Bibliografía"):
            with st.spinner("Buscando PDFs en todas las carpetas y moviéndolos a Bibliografía..."):
                new_df, logs = force_move_all_pdfs_to_bibliography()
                st.success("PDFs normalizados.")
                with st.expander("Ver log de movimiento de PDFs", expanded=True):
                    for line in logs:
                        st.write(line)
                st.rerun()

    inventory_df = load_inventory()

    if inventory_df.empty:
        st.info("Todavía no hay archivos guardados.")
        st.stop()

    st.dataframe(inventory_df, use_container_width=True)

    csv_bytes = inventory_df.to_csv(index=False).encode("utf-8-sig")

    st.download_button(
        "Descargar inventario CSV",
        data=csv_bytes,
        file_name="inventario_casa_matriz.csv",
        mime="text/csv",
    )

    st.stop()
